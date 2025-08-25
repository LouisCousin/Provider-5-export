"""Outils d'exportation des résultats de batch en document DOCX."""

from __future__ import annotations

import io
import json
from dataclasses import asdict, is_dataclass
from typing import Any, Dict, List

import markdown as md
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


class MarkdownToDocxConverter:
    """Convertit du texte Markdown en éléments DOCX."""

    def __init__(self, document: Document, styles: Dict[str, Dict[str, Any]]):
        """Initialise le convertisseur avec un document et un dictionnaire de styles."""

        self.doc = document
        self.styles = styles or {}

    def _apply_style(
        self,
        run,
        style_overrides: Dict[str, Any] | None = None,
        *,
        style_name: str = "response",
    ) -> None:
        """Applique un style au ``run`` donné.

        ``style_name`` permet de sélectionner un style de base dans ``self.styles``.
        ``style_overrides`` peut être utilisé pour modifier certains attributs.
        """

        style = {**self.styles.get(style_name, {}), **(style_overrides or {})}

        if font_name := style.get("font_name"):
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if size := style.get("font_size"):
            run.font.size = Pt(size)
        if color := style.get("font_color_rgb"):
            run.font.color.rgb = RGBColor(*color)
        run.bold = style.get("is_bold", False)
        run.italic = style.get("is_italic", False)

    def _add_inline(self, paragraph, node) -> None:
        """Ajoute récursivement les noeuds inline à un paragraphe."""

        if node.text:
            run = paragraph.add_run(node.text)
            self._apply_style(run)

        for child in list(node):
            if child.tag in {"strong", "b"}:
                run = paragraph.add_run(child.text or "")
                self._apply_style(run, {"is_bold": True})
            elif child.tag in {"em", "i"}:
                run = paragraph.add_run(child.text or "")
                self._apply_style(run, {"is_italic": True})
            elif child.tag == "code":
                run = paragraph.add_run(child.text or "")
                self._apply_style(run)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            elif child.tag == "a":
                text = child.text or ""
                href = child.get("href", "")
                run = paragraph.add_run(text)
                self._apply_style(run)
                if href and href != text:
                    tail_run = paragraph.add_run(f" ({href})")
                    self._apply_style(tail_run)
            else:
                self._add_inline(paragraph, child)

            if child.tail:
                tail = paragraph.add_run(child.tail)
                self._apply_style(tail)

    def _process_element(self, elem, list_style: str | None = None) -> None:
        """Traite les éléments HTML convertis depuis le Markdown."""

        tag = elem.tag
        if tag in {"p", "li"}:
            paragraph = (
                self.doc.add_paragraph(style=list_style)
                if list_style
                else self.doc.add_paragraph()
            )
            self._add_inline(paragraph, elem)
            for child in list(elem):
                if child.tag in {"ul", "ol"}:
                    self._process_element(
                        child,
                        "List Bullet" if child.tag == "ul" else "List Number",
                    )
        elif tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(tag[1])
            paragraph = self.doc.add_heading(level=level)
            self._add_inline(paragraph, elem)
        elif tag == "ul":
            for li in elem.findall("li"):
                self._process_element(li, "List Bullet")
        elif tag == "ol":
            for li in elem.findall("li"):
                self._process_element(li, "List Number")
        elif tag == "pre":
            code_text = "".join(elem.itertext()).strip()
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(code_text)
            self._apply_style(run)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        elif tag == "table":
            rows = elem.findall("tr")
            if rows:
                cols = len(rows[0].findall("th")) or len(rows[0].findall("td"))
                table = self.doc.add_table(rows=len(rows), cols=cols)
                for r_idx, row in enumerate(rows):
                    cells = row.findall("th") or row.findall("td")
                    for c_idx, cell in enumerate(cells):
                        paragraph = table.cell(r_idx, c_idx).paragraphs[0]
                        self._add_inline(paragraph, cell)
        else:
            if elem.text:
                paragraph = self.doc.add_paragraph()
                self._add_inline(paragraph, elem)

    def add_markdown(self, text: str) -> None:
        """Convertit un texte Markdown et l'ajoute au document."""

        if not text:
            return

        md_converter = md.Markdown(extensions=["fenced_code", "tables"])
        html = md_converter.convert(text)

        from xml.etree import ElementTree as ET

        root = ET.fromstring(f"<root>{html}</root>")
        for elem in list(root):
            self._process_element(elem)


def generer_export_docx(resultats: List[Any], styles: Dict[str, Dict[str, Any]]) -> io.BytesIO:
    """Génère un document DOCX à partir d'une liste de résultats de batch.

    Chaque résultat doit contenir au minimum les champs ``status``,
    ``prompt_text`` et ``clean_response`` (ou ``response``).
    """

    document = Document()
    converter = MarkdownToDocxConverter(document, styles)

    succeeded: List[Dict[str, Any]] = []
    failed: List[Dict[str, Any]] = []

    for res in resultats:
        data = asdict(res) if is_dataclass(res) else dict(res)
        if data.get("status") == "succeeded":
            succeeded.append(data)
        else:
            failed.append(data)

    # Section principale : prompts et réponses
    for item in succeeded:
        prompt_text = item.get("prompt_text", "")
        response_text = item.get("clean_response") or item.get("response", "")

        para = converter.doc.add_paragraph()
        run = para.add_run(prompt_text)
        converter._apply_style(run, style_name="prompt")

        converter.add_markdown(response_text)
        converter.doc.add_paragraph()

    # Section annexe pour les erreurs
    if failed:
        converter.doc.add_page_break()
        converter.doc.add_heading("Annexe - Requêtes échouées", level=1)
        for item in failed:
            title = item.get("prompt_text") or item.get("custom_id")
            converter.doc.add_paragraph(title, style="List Bullet")
            err = item.get("error")
            if isinstance(err, dict):
                err = json.dumps(err, ensure_ascii=False, indent=2)
            converter.doc.add_paragraph(str(err))

    output = io.BytesIO()
    converter.doc.save(output)
    output.seek(0)
    return output

